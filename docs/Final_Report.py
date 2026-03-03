from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ── Cover Page ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Cryptocurrency Volatility Prediction")
run.font.size = Pt(28)
run.font.bold = True

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Final Report")
run2.font.size = Pt(18)
run2.font.bold = True

doc.add_paragraph("Date: March 2026").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("Institution: PW Skills").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# ── 1. Executive Summary ──
h = doc.add_paragraph("1. Executive Summary")
h.runs[0].font.bold = True
h.runs[0].font.size = Pt(16)

doc.add_paragraph("This project builds a Machine Learning model to predict cryptocurrency volatility using historical Bitcoin market data. The XGBoost Regressor was trained on 23 engineered features and achieved 81.83% accuracy.")

# Results Table
table = doc.add_table(rows=6, cols=2)
table.style = 'Table Grid'
data = [["Metric","Value"],["R² Score","0.8183"],["RMSE","0.007201"],["MAE","0.004658"],["Training Rows","3,012"],["Features Used","23"]]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        cell = table.rows[i].cells[j]
        cell.text = val
        if i == 0:
            cell.paragraphs[0].runs[0].font.bold = True
doc.add_page_break()

# ── 2. Problem Statement ──
h = doc.add_paragraph("2. Problem Statement")
h.runs[0].font.bold = True
h.runs[0].font.size = Pt(16)
doc.add_paragraph("Cryptocurrency markets are highly volatile. Volatility refers to the degree of variation in the price of a cryptocurrency over time. High volatility leads to significant risks for traders and investors.")
for t in ["Build ML model to predict cryptocurrency volatility","Use OHLC prices, trading volume, market cap","Anticipate periods of heightened volatility","Help traders manage risk and make informed decisions"]:
    doc.add_paragraph(t, style='List Bullet')
doc.add_page_break()

# ── 3. Dataset Description ──
h = doc.add_paragraph("3. Dataset Description")
h.runs[0].font.bold = True
h.runs[0].font.size = Pt(16)
table = doc.add_table(rows=9, cols=3)
table.style = 'Table Grid'
data = [["Column","Type","Description"],["date","object","Trading date"],["crypto_name","object","Cryptocurrency name"],["open","float64","Opening price (USD)"],["high","float64","Highest price of day"],["low","float64","Lowest price of day"],["close","float64","Closing price (USD)"],["volume","float64","Trading volume"],["marketCap","float64","Market capitalization"]]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        cell = table.rows[i].cells[j]
        cell.text = val
        if i == 0:
            cell.paragraphs[0].runs[0].font.bold = True
for t in ["Total Records: 72,946","Cryptocurrencies: 50+","Missing Values: 0","Focus: Bitcoin (3,248 rows)","Date Range: 2013 to 2023"]:
    doc.add_paragraph(t, style='List Bullet')
doc.add_page_break()

# ── 4. Methodology ──
h = doc.add_paragraph("4. Methodology")
h.runs[0].font.bold = True
h.runs[0].font.size = Pt(16)
table = doc.add_table(rows=7, cols=3)
table.style = 'Table Grid'
data = [["Step","Action","Output"],["1","Data Collection","72,946 rows loaded, Bitcoin filtered"],["2","Preprocessing","Cleaned data, bitcoin_clean.csv"],["3","EDA","Price trends, volatility graphs, heatmap"],["4","Feature Engineering","23 features, bitcoin_features.csv"],["5","Model Training","Trained RF and XGBoost, compared"],["6","Deployment","Streamlit app at localhost:8501"]]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        cell = table.rows[i].cells[j]
        cell.text = val
        if i == 0:
            cell.paragraphs[0].runs[0].font.bold = True
doc.add_page_break()

# ── 5. Feature Engineering ──
h = doc.add_paragraph("5. Feature Engineering Summary")
h.runs[0].font.bold = True
h.runs[0].font.size = Pt(16)
table = doc.add_table(rows=8, cols=3)
table.style = 'Table Grid'
data = [["Category","Features","Count"],["Target","volatility_7d","1"],["Price","MA_7, MA_14, MA_30, price_range, price_range_pct, close_open_pct","6"],["Technical","BB_width, ATR_14, RSI_14","3"],["Volume","volume_MA7, volume_change, volume_spike, liquidity_ratio","4"],["Lag","volatility_lag1/3/7, return_lag1/3","5"],["OHLCV Base","open, high, low, close, volume","5"],["TOTAL","","23"]]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        cell = table.rows[i].cells[j]
        cell.text = val
        if i == 0 or i == 7:
            cell.paragraphs[0].runs[0].font.bold = True
doc.add_page_break()

# ── 6. Model Results ──
h = doc.add_paragraph("6. Model Results")
h.runs[0].font.bold = True
h.runs[0].font.size = Pt(16)
table = doc.add_table(rows=3, cols=5)
table.style = 'Table Grid'
data = [["Model","RMSE","MAE","R² Score","Winner"],["Random Forest","0.007467","0.005019","0.8046",""],["XGBoost","0.007201","0.004658","0.8183","✅ Best"]]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        cell = table.rows[i].cells[j]
        cell.text = val
        if i == 0:
            cell.paragraphs[0].runs[0].font.bold = True
for t in ["R² = 0.8183 means model is 81.83% accurate","RMSE = 0.007 — very low prediction error","MAE = 0.004 — average error is only 0.004","XGBoost won on all 3 metrics"]:
    doc.add_paragraph(t, style='List Bullet')
doc.add_page_break()

# ── 7. Key Insights ──
h = doc.add_paragraph("7. Key Insights")
h.runs[0].font.bold = True
h.runs[0].font.size = Pt(16)
insights = [("Past Volatility Predicts Future","volatility_lag1 was most important feature (78%). Volatility follows strong temporal patterns."),("Price Range Matters","price_range_pct was 2nd most important. Large daily price swings indicate volatile periods."),("Volume Spikes Signal Volatility","Unusual volume spikes often precede volatile price movements."),("Bollinger Bands Work","BB_width was among top features. Wider bands mean higher expected volatility."),("Bitcoin Has Volatile Phases","Bitcoin showed extreme volatility during 2017-2018 and 2020-2021 bull runs.")]
for title, desc in insights:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.bold = True
    run.font.size = Pt(12)
    doc.add_paragraph(desc)
doc.add_page_break()

# ── 8. Deployment Summary ──
h = doc.add_paragraph("8. Deployment Summary")
h.runs[0].font.bold = True
h.runs[0].font.size = Pt(16)
doc.add_paragraph("Tool: Streamlit  |  URL: localhost:8501")
table = doc.add_table(rows=4, cols=3)
table.style = 'Table Grid'
data = [["Risk Level","Condition","Meaning"],["HIGH","Volatility > 0.05","Market is highly unstable"],["MEDIUM","0.02 to 0.05","Proceed with caution"],["LOW","Volatility < 0.02","Market is stable"]]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        cell = table.rows[i].cells[j]
        cell.text = val
        if i == 0:
            cell.paragraphs[0].runs[0].font.bold = True
doc.add_page_break()

# ── 9. Conclusion ──
h = doc.add_paragraph("9. Conclusion")
h.runs[0].font.bold = True
h.runs[0].font.size = Pt(16)
doc.add_paragraph("This project successfully achieved all objectives. A production-ready ML pipeline was built from raw data to deployed web application.")
p = doc.add_paragraph()
p.add_run("Key Achievements:").font.bold = True
for t in ["Processed 72,946 cryptocurrency records","Engineered 23 meaningful features","Trained and compared 2 ML models","XGBoost achieved R² = 0.8183 (81% accurate)","Deployed as Streamlit web application","Real-time prediction with color-coded risk levels"]:
    doc.add_paragraph(t, style='List Bullet')
p2 = doc.add_paragraph()
p2.add_run("Future Improvements:").font.bold = True
for t in ["Add more cryptocurrencies (ETH, BNB, SOL)","Use LSTM deep learning for better accuracy","Add news sentiment analysis as feature","Deploy on cloud (AWS / Streamlit Cloud)","Add live price API (CoinGecko API)"]:
    doc.add_paragraph(t, style='List Bullet')

# ── Save ──
doc.save('Final_Report.docx')
print("✅ Final_Report.docx saved!")